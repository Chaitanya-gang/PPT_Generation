"""
newd2p - DOCX Parser
"""

from pathlib import Path
from docx import Document
from src.parsers.base_parser import BaseParser
from src.parsers.models import ParsedDocument, DocumentSection, ExtractedTable
from src.utils.text_cleaner import clean_text
from src.utils.logger import get_logger

logger = get_logger("docx_parser")


class DocxParser(BaseParser):

    def can_parse(self, file_extension: str) -> bool:
        return file_extension.lower() in [".docx", ".doc"]

    def parse(self, file_path: str, file_id: str) -> ParsedDocument:
        logger.info(f"Parsing DOCX: {file_path}")

        path = Path(file_path)
        doc = Document(file_path)

        raw_text = self._extract_text(doc)
        cleaned = clean_text(raw_text)
        sections = self._extract_sections(doc)
        tables = self._extract_tables(doc)

        total_pages = max(1, len(raw_text) // 3000)

        parsed = ParsedDocument(
            file_id=file_id,
            filename=path.name,
            file_type="docx",
            total_pages=total_pages,
            raw_text=raw_text,
            cleaned_text=cleaned,
            sections=sections,
            tables=tables,
            metadata={
                "file_size": path.stat().st_size,
                "paragraphs": len(doc.paragraphs),
                "tables_count": len(doc.tables),
            }
        )

        logger.info(f"DOCX parsed: {parsed.total_words} words, {len(sections)} sections, {len(tables)} tables")
        return parsed

    def _extract_text(self, doc: Document) -> str:
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return "\n\n".join(text_parts)

    def _extract_sections(self, doc: Document) -> list:
        sections = []
        current_section = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if heading style
            is_heading = False
            level = 1

            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if "heading" in style_name:
                    is_heading = True
                    # Extract heading level
                    for char in style_name:
                        if char.isdigit():
                            level = int(char)
                            break
                elif "title" in style_name:
                    is_heading = True
                    level = 1

            # Also detect bold short lines as headings
            if not is_heading and len(text) < 100:
                runs = para.runs
                if runs and all(run.bold for run in runs if run.text.strip()):
                    is_heading = True
                    level = 2

            if is_heading:
                # Save previous section
                if current_section:
                    sections.append(DocumentSection(
                        level=current_section["level"],
                        title=current_section["title"],
                        content="\n".join(current_content).strip(),
                    ))
                current_section = {"title": text, "level": level}
                current_content = []
            else:
                current_content.append(text)

        # Save last section
        if current_section:
            sections.append(DocumentSection(
                level=current_section["level"],
                title=current_section["title"],
                content="\n".join(current_content).strip(),
            ))
        elif current_content:
            sections.append(DocumentSection(
                level=1,
                title="Main Content",
                content="\n".join(current_content).strip(),
            ))

        return sections

    def _extract_tables(self, doc: Document) -> list:
        tables = []

        for idx, table in enumerate(doc.tables):
            try:
                rows_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows_data.append(row_data)

                if len(rows_data) < 2:
                    continue

                headers = rows_data[0]
                data_rows = rows_data[1:]

                tables.append(ExtractedTable(
                    table_id=f"table_{idx+1}",
                    headers=headers,
                    rows=data_rows,
                ))

            except Exception as e:
                logger.error(f"Error extracting table {idx}: {e}")

        return tables
